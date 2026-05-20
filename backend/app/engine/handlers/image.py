import os
import shutil
import logging
from pathlib import Path
from PIL import Image
from PIL import ImageOps
from PIL import ImageFilter
import pytesseract
from pytesseract import Output
from docx import Document
from docx.shared import Inches
from app.engine.handlers.base import BaseHandler
try:
    import pillow_heif
    pillow_heif.register_heif_opener()
    HEIF_PLUGIN_ENABLED = True
except Exception:
    HEIF_PLUGIN_ENABLED = False
try:
    import easyocr
    EASYOCR_ENABLED = True
except Exception:
    EASYOCR_ENABLED = False

logger = logging.getLogger(__name__)

class ImageHandler(BaseHandler):
    """
    Handles all image processing and conversion pathways.
    Supports standard Pillow formats, HEIC/AVIF via ImageMagick & libheif,
    direct Image-to-PDF rendering, and Tesseract-powered Image-to-DOCX OCR.
    """

    def convert(self, input_path: Path, output_path: Path, from_ext: str, to_ext: str, quality: str = "fast") -> Path:
        from_ext = from_ext.lower()
        to_ext = to_ext.lower()

        logger.info(f"ImageHandler converting {from_ext} -> {to_ext}")

        # 1. OCR Image to DOCX
        if to_ext == ".docx":
            return self._to_docx(input_path, output_path)

        # 2. Image to PDF
        if to_ext == ".pdf":
            return self._to_pdf(input_path, output_path, from_ext)

        # 3. Image-to-Image conversions (including HEIC, AVIF, GIF, JPG, PNG, WEBP)
        return self._to_image(input_path, output_path, from_ext, to_ext)

    def _to_pdf(self, input_path: Path, output_path: Path, from_ext: str) -> Path:
        """
        Converts an image directly to a high-fidelity PDF page.
        """
        temp_png = None
        try:
            img = self._load_image(input_path)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            img.save(output_path, "PDF")
            logger.info(f"Successfully converted image to PDF: {output_path}")
            return output_path
        finally:
            if temp_png and temp_png.exists():
                temp_png.unlink()

    def _to_docx(self, input_path: Path, output_path: Path) -> Path:
        """
        Converts an image to an editable Word Document using Tesseract OCR.
        Falls back to embedding the original image if OCR confidence is low.
        """
        temp_png = None
        try:
            logger.info("Running Tesseract OCR on image...")
            img = self._load_image(input_path)
            ocr_img = self._preprocess_for_ocr(img)
            
            # Extract text and geometric OCR blocks for better line reconstruction.
            try:
                extracted_text = pytesseract.image_to_string(ocr_img, timeout=45)
                data = pytesseract.image_to_data(ocr_img, output_type=Output.DICT, timeout=45)
            except Exception as e:
                logger.warning(f"Tesseract extraction timed out or failed: {e}. Falling back to image embedding.")
                extracted_text = ""
                data = None

            cleaned_text = extracted_text.strip()
            if len(cleaned_text) <= 15:
                easy_text = self._extract_text_easyocr(ocr_img)
                if easy_text:
                    cleaned_text = easy_text
            
            doc = Document()
            # If we found sufficient text, write editable DOCX
            if len(cleaned_text) > 15:
                logger.info(f"Sufficient text detected ({len(cleaned_text)} chars). Generating editable paragraphs...")
                doc.add_paragraph(f"--- OCR Extracted Text from: {input_path.name} ---")
                doc.add_paragraph("")
                structured_lines = self._lines_from_ocr_data(data)
                if structured_lines:
                    for line in structured_lines:
                        doc.add_paragraph(line)
                else:
                    for line in cleaned_text.split("\n"):
                        line_strip = line.strip()
                        if line_strip:
                            doc.add_paragraph(line_strip)
            else:
                # Scanned confidence low, embed picture directly
                logger.info("Low OCR confidence/length. Embedding picture as high-fidelity fallback.")
                doc.add_paragraph(f"--- Embedded Image: {input_path.name} ---")
                doc.add_paragraph("")
                embed_path = input_path
                if input_path.suffix.lower() == ".heic":
                    temp_png = input_path.parent / f"{input_path.stem}_embed_temp.png"
                    img.convert("RGB").save(temp_png)
                    embed_path = temp_png
                doc.add_picture(str(embed_path), width=Inches(6.0))

            doc.save(output_path)
            logger.info(f"DOCX created successfully: {output_path}")
            return output_path

        finally:
            if temp_png and temp_png.exists():
                temp_png.unlink()

    def _lines_from_ocr_data(self, data) -> list[str]:
        if not data or "text" not in data:
            return []

        grouped: dict[tuple[int, int, int], list[tuple[int, str]]] = {}
        for i, raw_text in enumerate(data.get("text", [])):
            text = (raw_text or "").strip()
            if not text:
                continue
            conf_raw = str(data.get("conf", ["-1"])[i])
            try:
                conf = float(conf_raw)
            except ValueError:
                conf = -1.0
            if conf < 45:
                continue
            key = (
                int(data.get("block_num", [0])[i]),
                int(data.get("par_num", [0])[i]),
                int(data.get("line_num", [0])[i]),
            )
            left = int(data.get("left", [0])[i])
            grouped.setdefault(key, []).append((left, text))

        lines = []
        for key in sorted(grouped.keys()):
            words = [w for _, w in sorted(grouped[key], key=lambda it: it[0])]
            line = " ".join(words).strip()
            if line:
                lines.append(line)
        return lines

    def _to_image(self, input_path: Path, output_path: Path, from_ext: str, to_ext: str) -> Path:
        """
        Transforms images using ImageMagick as first-line driver,
        falling back to Pillow + libheif CLI encoding.
        """
        # Primary Pipeline: ImageMagick (handles AVIF, HEIC, and other formats natively)
        try:
            cmd = ["convert", str(input_path), str(output_path)]
            logger.info(f"Running ImageMagick: {' '.join(cmd)}")
            self.run_subprocess(cmd, timeout=30)
            if output_path.exists():
                return output_path
        except Exception as e:
            logger.warning(f"ImageMagick failed: {e}. Swapping to Pillow / libheif fallback...")

        # Secondary Pipeline: Pillow + libheif wrappers
        temp_png = None
        try:
            # 1. Decode HEIC if it's the source
            if from_ext == ".heic":
                source_img = input_path
            else:
                source_img = input_path

            # 2. Encode to HEIC if it's the target
            if to_ext == ".heic":
                img = self._load_image(source_img)
                try:
                    cmd = ["heif-enc", str(source_img), str(output_path)]
                    self.run_subprocess(cmd, timeout=30)
                    return output_path
                except Exception:
                    return self._save_as_heic(img, output_path)

            # 3. Handle standard Pillow transformations (JPG, PNG, WEBP, GIF, AVIF)
            img = self._load_image(source_img)
            
            # Format modifications for transparency loss in JPEGs
            if to_ext in [".jpg", ".jpeg"] and img.mode != "RGB":
                img = img.convert("RGB")
            save_kwargs = self._metadata_kwargs(img)
            img.save(output_path, **save_kwargs)
            return output_path

        finally:
            if temp_png and temp_png.exists():
                temp_png.unlink()
    def _load_image(self, path: Path) -> Image.Image:
        try:
            return Image.open(path)
        except Exception:
            if path.suffix.lower() == ".heic":
                temp_png = path.parent / f"{path.stem}_decode_fallback.png"
                cmd = ["heif-dec", str(path), str(temp_png)]
                self.run_subprocess(cmd, timeout=30)
                return Image.open(temp_png)
            raise

    def _save_as_heic(self, img: Image.Image, output_path: Path) -> Path:
        if HEIF_PLUGIN_ENABLED:
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(output_path, format="HEIF", quality=90)
            return output_path
        raise RuntimeError("No HEIC encoder available in runtime.")

    def _metadata_kwargs(self, img: Image.Image) -> dict:
        kwargs = {}
        exif = img.info.get("exif")
        icc_profile = img.info.get("icc_profile")
        if exif:
            kwargs["exif"] = exif
        if icc_profile:
            kwargs["icc_profile"] = icc_profile
        return kwargs

    def _extract_text_easyocr(self, img: Image.Image) -> str:
        if not EASYOCR_ENABLED:
            return ""

    def _preprocess_for_ocr(self, img: Image.Image) -> Image.Image:
        # Normalize color and contrast to improve OCR accuracy, especially for HEIC captures.
        base = img.convert("L")
        boosted = ImageOps.autocontrast(base, cutoff=2)
        denoised = boosted.filter(ImageFilter.MedianFilter(size=3))
        return denoised
        try:
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)
            results = reader.readtext(img, detail=1, paragraph=False)
            if not results:
                return ""
            filtered = []
            for _, text, conf in results:
                if text and conf >= 0.35:
                    filtered.append(text.strip())
            return "\n".join([t for t in filtered if t])
        except Exception as e:
            logger.warning(f"EasyOCR fallback failed: {e}")
            return ""
